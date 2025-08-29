import io
import re
import csv
import zipfile
from datetime import datetime
from urllib.parse import urlparse
from zoneinfo import ZoneInfo
from pathlib import Path

import streamlit as st
import requests
from bs4 import BeautifulSoup, Tag, NavigableString, Comment, Doctype, ProcessingInstruction
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

# -------------------------
# CONFIG / CONSTANTS
# -------------------------
ALWAYS_STRIP = {"script", "style", "noscript", "template"}
INLINE_TAGS = {"a","span","strong","em","b","i","u","s","small","sup","sub","mark","abbr","time","code","var","kbd"}
DEFAULT_EXCLUDE = [
    "header", "footer", "nav",
    ".cookie", ".newsletter",
    "[class*='breadcrumb']",
    "[class*='wishlist']",
    "[class*='simplesearch']",
    "[id*='gallery']",
    "[class*='usp']",
    "[class*='feefo']",
    "[class*='associated-blogs']",
    "[class*='popular']",
    # Explore/SPA results containers and variants
    ".sr-main.js-searchpage-content.visible",
    "[class~='sr-main'][class~='js-searchpage-content'][class~='visible']",
    "[class*='js-searchpage-content']",
    "[class*='searchpage-content']",
    # Map modal container to exclude
    ".lmd-map-modal-create.js-lmd-map-modal-map",
]
DATE_TZ = "Europe/London"
DATE_FMT = "%d/%m/%Y"

NOISE_SUBSTRINGS = (
    "google tag manager",
    "loading results",
    "load more",
    "updating results",
    "something went wrong",
    "filters",
    "apply filters",
    "clear",
    "sort by",
    "to collect end-user usage analytics",
    "place this code immediately before the closing",
)

# -------------------------
# UTILITIES
# -------------------------

from pathlib import Path
import streamlit as st

# ---- Page config MUST be the first Streamlit call ----
APP_DIR = Path(__file__).resolve().parent
ICON_CANDIDATES = [
    APP_DIR / "assets" / "JAFavicon.png",
    APP_DIR / "JAFavicon.png",
]

icon_path = next((p for p in ICON_CANDIDATES if p.exists()), None)

st.set_page_config(
    page_title="Content Rec Template Tool",
    page_icon=str(icon_path) if icon_path else "🧩",
    layout="wide",
)

if not icon_path:
    st.sidebar.warning("Favicon not found (looked in assets/JAFavicon.png and JAFavicon.png).")

def uk_today_str() -> str:
    return datetime.now(ZoneInfo(DATE_TZ)).strftime(DATE_FMT)

def clean_slug_to_name(slug: str) -> str:
    return slug.replace("-", " ").strip().title()

def fallback_page_name_from_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    parts = [p for p in path.split("/") if p]
    try:
        i = parts.index("destinations")
        if len(parts) > i + 2:
            return clean_slug_to_name(parts[i + 2])
    except ValueError:
        pass
    return clean_slug_to_name(parts[-1] if parts else (urlparse(url).hostname or "Page"))

@st.cache_data(show_spinner=False, ttl=3600)
def fetch_html(url: str) -> tuple[str, bytes]:
    resp = requests.get(
        url,
        timeout=30,
        headers={"User-Agent": "Mozilla/5.0 (compatible; ContentRecTool/1.0)"},
    )
    resp.raise_for_status()
    return resp.url, resp.content

def normalise_keep_newlines(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"[ \t]*\n[ \t]*", "\n", s)
    return s

def is_noise(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    return any(sub in t for sub in NOISE_SUBSTRINGS)

def annotate_anchor_text(a: Tag, annotate_links: bool) -> str:
    text = a.get_text(" ", strip=True)
    href = a.get("href", "")
    return f"{text} (→ {href})" if (annotate_links and href) else text

def extract_text_preserve_breaks(node: Tag | NavigableString, annotate_links: bool) -> str:
    """Extract visible text; convert <br> to \n; handle anchors as one unit."""
    if isinstance(node, NavigableString):
        return str(node)
    parts = []
    for child in node.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif isinstance(child, Tag):
            if child.name == "br":
                parts.append("\n")
            elif child.name == "a":
                parts.append(annotate_anchor_text(child, annotate_links))
            else:
                parts.append(extract_text_preserve_breaks(child, annotate_links))
    return "".join(parts)

def extract_signposted_lines_from_body(body: Tag, annotate_links: bool, include_img_src: bool = False) -> list[str]:
    """
    Emit ONLY:
      - <h1> … <h6> lines
      - <p> lines
      - <img alt="…"> (or <img alt="…" src="…"> when enabled) for every <img> encountered

    Lists are flattened to <p>. Critically, <p> is split on <br> and blank lines preserved
    (blank <p> emitted as '<p>' with no text).

    Additionally, capture stray text nodes (bare text in containers) as <p>, but skip
    comments/doctype/processing instructions and obvious UI/analytics noise.
    """
    lines: list[str] = []

    def emit_lines(tag_name: str, text: str):
        text = normalise_keep_newlines(text)
        segments = text.split("\n")
        for seg in segments:
            seg_stripped = seg.strip()
            if seg_stripped:
                if tag_name == "p" and is_noise(seg_stripped):
                    continue
                lines.append(f"<{tag_name}> {seg_stripped}")
            else:
                if tag_name == "p":
                    lines.append("<p>")

    def emit_img(img_tag: Tag):
        if not isinstance(img_tag, Tag) or img_tag.name != "img":
            return
        alt = (img_tag.get("alt") or "").strip().replace('"', '\\"')
        if include_img_src:
            src = (img_tag.get("src") or "").strip().replace('"', '\\"')
            if src:
                lines.append(f'<img alt="{alt}" src="{src}">')
                return
        lines.append(f'<img alt="{alt}">')

    def handle(tag: Tag):
        name = tag.name
        if name in ALWAYS_STRIP:
            return

        # Headings
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            txt = extract_text_preserve_breaks(tag, annotate_links)
            if txt.strip():
                emit_lines(name, txt)
            return

        # Paragraphs
        if name == "p":
            txt = tag.get_text(" ", strip=True)
            if txt.strip():
                emit_lines("p", txt)
            for img in tag.find_all("img"):
                emit_img(img)
            return

        # Lists
        if name in {"ul", "ol"}:
            for li in tag.find_all("li", recursive=False):
                txt = extract_text_preserve_breaks(li, annotate_links)
                if txt.strip():
                    emit_lines("p", txt)
                for img in li.find_all("img"):
                    emit_img(img)
                for sub in li.find_all(["ul", "ol"], recursive=False):
                    for sub_li in sub.find_all("li", recursive=False):
                        sub_txt = extract_text_preserve_breaks(sub_li, annotate_links)
                        if sub_txt.strip():
                            emit_lines("p", sub_txt)
                        for img in sub_li.find_all("img"):
                            emit_img(img)
            return

        # Generic containers: group contiguous inline content; recurse into block-level children
        buf = []
        def flush_buf():
            if not buf:
                return
            joined = normalise_keep_newlines("".join(buf))
            if joined.strip() and not is_noise(joined):
                emit_lines("p", joined)
            buf.clear()

        for child in tag.children:
            if isinstance(child, (Comment, Doctype, ProcessingInstruction)):
                continue
            if isinstance(child, NavigableString):
                buf.append(str(child))
            elif isinstance(child, Tag):
                if child.name == "br":
                    buf.append("\n")
                elif child.name == "img":
                    flush_buf()
                    emit_img(child)
                elif child.name in INLINE_TAGS:
                    buf.append(extract_text_preserve_breaks(child, annotate_links))
                else:
                    flush_buf()
                    handle(child)
        flush_buf()

    for child in body.children:
        if isinstance(child, (Comment, Doctype, ProcessingInstruction)):
            continue
        if isinstance(child, NavigableString):
            raw = normalise_keep_newlines(str(child))
            if raw.strip() and not is_noise(raw):
                emit_lines("p", raw)
        elif isinstance(child, Tag):
            if child.name == "img":
                emit_img(child)
            else:
                handle(child)

    # Deduplicate trivial adjacent repeats
    deduped, prev = [], None
    for ln in lines:
        if ln != prev:
            deduped.append(ln)
        prev = ln
    return deduped

# -------------------------
# FIX: remove everything before first <h1> (all ancestor levels)
# -------------------------

def remove_before_first_h1_all_levels(body: Tag) -> None:
    """
    Remove *all* nodes that appear before the first <h1> in document order.
    Walks the ancestor chain from <body> down to <h1>; at each level
    removes previous siblings of the node on the path to the <h1>.
    """
    if body is None:
        return
    first_h1 = body.find("h1")
    if first_h1 is None:
        return

    # Build chain from the direct child of body to the h1
    chain = []
    node = first_h1
    while node is not None and node != body:
        chain.append(node)
        node = node.parent
    chain.reverse()

    for child in chain:
        for prev in list(child.previous_siblings):
            try:
                if isinstance(prev, Tag):
                    prev.decompose()
                elif isinstance(prev, NavigableString):
                    prev.extract()
            except Exception:
                continue

# -------------------------
# DOCX helpers
# -------------------------

def iter_paragraphs_and_tables(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def replace_placeholders_safe(doc: Document, mapping: dict[str, str]):
    keys = sorted(mapping.keys(), key=len, reverse=True)
    for p in iter_paragraphs_and_tables(doc):
        t = p.text or ""
        replaced = False
        for k in keys:
            v = mapping[k]
            if k in t:
                t = t.replace(k, v)
                replaced = True
        if replaced:
            for r in list(p.runs):
                r.clear()
            p.clear()
            p.add_run(t)

def find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for p in iter_paragraphs_and_tables(doc):
        if placeholder in (p.text or ""):
            return p
    return None

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def replace_placeholder_with_lines(doc: Document, placeholder: str, lines: list[str]):
    target = find_placeholder_paragraph(doc, placeholder)
    if target is None:
        raise ValueError(f"Placeholder '{placeholder}' not found in template.")
    if not lines:
        target.clear()
        return
    target.clear()
    target.add_run(lines[0])
    anchor = target
    for line in lines[1:]:
        anchor = insert_paragraph_after(anchor, line)

def build_docx(template_bytes: bytes, meta: dict, lines: list[str]) -> bytes:
    bio = io.BytesIO(template_bytes)
    doc = Document(bio)
    replace_placeholders_safe(doc, {
        "[PAGE]": meta.get("page", ""),
        "[DATE]": meta.get("date", ""),
        "[URL]": meta.get("url", ""),
        "[TITLE]": meta.get("title", ""),
        "[TITLE LENGTH]": str(meta.get("title_len", 0)),
        "[DESCRIPTION]": meta.get("description", ""),
        "DESCRIPTION": meta.get("description", ""),
        "[DESCRIPTION LENGTH]": str(meta.get("description_len", 0)),
        "[AGENCY]": meta.get("agency", ""),
        "[CLIENT NAME]": meta.get("client_name", ""),
    })
    replace_placeholder_with_lines(doc, "[PAGE BODY CONTENT]", lines)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# -------------------------
# CORE PROCESS
# -------------------------

def first_h1_text(soup: BeautifulSoup) -> str | None:
    if not soup.body:
        return None
    h1 = soup.body.find("h1")
    if not h1:
        return None
    txt = extract_text_preserve_breaks(h1, annotate_links=False)
    txt = normalise_keep_newlines(txt)
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip() or None

def process_url(
    url: str,
    exclude_selectors: list[str],
    annotate_links: bool = False,
    remove_before_h1: bool = False,
    include_img_src: bool = False,
):
    final_url, html_bytes = fetch_html(url)
    soup = BeautifulSoup(html_bytes, "lxml")

    # global strip (script/style/noscript/template)
    for el in soup.find_all(list(ALWAYS_STRIP)):
        el.decompose()

    body = soup.body or soup

    # exclude universal blocks
    for sel in exclude_selectors:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass

    # hard-kill: ensure any element with all three classes is removed even if selector order changes
    try:
        for el in body.find_all(lambda t: isinstance(t, Tag) and t.has_attr('class') and {'sr-main','js-searchpage-content','visible'}.issubset(set(t.get('class', [])))):
            el.decompose()
    except Exception:
        pass

    # Also explicitly remove via robust CSS selectors (belt-and-braces)
    for sel in [
        '.sr-main.js-searchpage-content.visible',
        "[class~='sr-main'][class~='js-searchpage-content'][class~='visible']",
        "[class*='js-searchpage-content']",
        "[class*='searchpage-content']",
        ".lmd-map-modal-create.js-lmd-map-modal-map",
    ]:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass

    # If requested, remove everything before the first <h1> (robust across ancestor levels)
    if remove_before_h1:
        remove_before_first_h1_all_levels(body)

    # extract signposted lines
    lines = extract_signposted_lines_from_body(body, annotate_links=annotate_links, include_img_src=include_img_src)

    # meta
    head = soup.head or soup
    title = head.title.string.strip() if (head and head.title and head.title.string) else "N/A"
    meta_el = head.find("meta", attrs={"name": "description"}) if head else None
    description = meta_el.get("content").strip() if (meta_el and meta_el.get("content")) else "N/A"

    # page name: prefer H1
    page_name = first_h1_text(soup) or fallback_page_name_from_url(final_url)

    meta = {
        "page": page_name,
        "date": uk_today_str(),
        "url": final_url,
        "title": title,
        "title_len": len(title) if title != "N/A" else 0,
        "description": description,
        "description_len": len(description) if description != "N/A" else 0,
    }
    return meta, lines

# -------------------------
# FILENAME SAFETY
# -------------------------

def safe_filename(name: str, maxlen: int = 120) -> str:
    # collapse any whitespace/newlines to single spaces
    name = re.sub(r"\s+", " ", name)
    # remove characters that break downloads
    name = re.sub(r'[\\/*?:"<>|]+', "", name)
    # commas are legal but can confuse some agents – make safer
    name = name.replace(",", "")
    # trim length and trailing dots/spaces
    return (name[:maxlen]).rstrip(". ")

# -------------------------
# STREAMLIT APP
# -------------------------

icon_path = "JAFavicon.png"
st.set_page_config(
    page_title="Content Rec Template Tool",
    page_icon=icon_path if Path(icon_path).exists() else None,
    layout="wide",
)

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700&display=swap');

/* Global font */
html, body, [data-testid="stAppViewContainer"] * { font-family: 'Montserrat', sans-serif; }

/* Hide Streamlit's Material icon spans to prevent 'keyboard_arrow_down' text overlap */
[data-testid="stIconMaterial"] { display: none !important; }

/* Main title */
section[tabindex="0"] h1:first-of-type {
  text-align: center;
  color: #4A90E2;
  font-size: 3em;
  padding-bottom: .5em;
  border-bottom: 2px solid #4A90E2;
}

/* Sidebar look + width */
[data-testid="stSidebar"] {
  background-color: #1a1e24;
  border-right: 1px solid #4A90E2;
  min-width: 320px;
  max-width: 420px;
}

/* Expander headers */
[data-testid="stExpander"] [data-testid="stExpanderHeader"] {
  background-color: #363945;
  border-radius: 8px;
  padding: 10px 15px;
  margin-bottom: 10px;
  border: none;
  font-weight: bold;
  color: #E0E0E0;
}

/* Buttons */
.stButton > button {
  width: 100%;
  background-color: #323640;
  color: #E0E0E0;
  border: 1px solid #4A90E2;
  border-radius: 8px;
  padding: 10px;
  transition: background-color .3s, color .3s;
}
.stButton > button:hover {
  background-color: #4A90E2;
  color: #fff;
  border-color: #fff;
}

/* Tabs */
[data-testid="stTabs"] button[role="tab"] { background-color: #323640; color: #E0E0E0; }
[data-testid="stTabs"] button[role="tab"][aria-selected="true"] {
  color: #4A90E2;
  box-shadow: inset 0 -3px 0 0 #4A90E2;
}
</style>
""",
    unsafe_allow_html=True,
)

st.title("Content Rec Template Generation Tool")

# session state for stable downloads across reruns
if "single_docx" not in st.session_state:
    st.session_state.single_docx = None
    st.session_state.single_docx_name = None
if "batch_zip" not in st.session_state:
    st.session_state.batch_zip = None
    st.session_state.batch_zip_name = "content_recommendations.zip"

with st.sidebar:
    st.header("Template & Options")
    tpl_file = st.file_uploader("Upload Template as .DOCX file", type=["docx"])
    st.caption("This should be your blank template with placeholders (e.g., [PAGE], [DATE], [PAGE BODY CONTENT], etc.).")

st.divider()
    st.subheader("Need a template?")

    from pathlib import Path
    APP_DIR = Path(__file__).resolve().parent
    TEMPLATE_CANDIDATES = [
        APP_DIR / "assets" / "blank_template.docx",
        APP_DIR / "blank_template.docx",
]

template_path = next((p for p in TEMPLATE_CANDIDATES if p.exists()), None)

if template_path:
    with open(template_path, "rb") as file:
        st.download_button(
            label="Download a Blank Template",
            data=file,
            file_name="blank_template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
else:
    st.info("Place your template at assets/blank_template.docx (preferred) or alongside app.py as blank_template.docx to enable this download.")

    st.divider()
    st.subheader("Exclude Selectors")
    exclude_txt = st.text_area(
        "Comma-separated CSS selectors to remove from <body>",
        value=", ".join(DEFAULT_EXCLUDE),
        height=120,
    )
    exclude_selectors = [s.strip() for s in exclude_txt.split(",") if s.strip()]

    st.subheader("Link formatting")
    annotate_links = st.toggle("Append (→ URL) after anchor text", value=False)

    # Toggle instead of checkbox for consistency with Link formatting
    remove_before_h1 = st.toggle("Delete everything before first <h1>", value=False)

    # Include <img> src alongside alt text
    include_img_src = st.toggle("Include <img> src in output", value=False)

    st.caption("Timezone fixed to Europe/London; dates in DD/MM/YYYY.")

# --- Tabs ---
tab1, tab2 = st.tabs(["Single URL", "Batch (CSV)"])

with tab1:
    st.subheader("Single page")

    # Agency / Client fields just above the URL field
    col0a, col0b = st.columns([1, 1])
    with col0a:
        agency_name = st.text_input("Agency Name", value="", placeholder="e.g., JA Consulting")
    with col0b:
        client_name = st.text_input("Client Name", value="", placeholder="e.g., Workspace")

    url = st.text_input("URL", value="https://www.example.com")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        do_preview = st.button("Extract preview")
    with col_b:
        do_doc = st.button("Generate DOCX")

    if do_preview or do_doc:
        if not tpl_file and do_doc:
            st.error("Please upload your Rec Template.docx in the sidebar first.")
        else:
            try:
                meta, lines = process_url(
                    url,
                    exclude_selectors,
                    annotate_links=annotate_links,
                    remove_before_h1=remove_before_h1,
                    include_img_src=include_img_src,
                )
                # Inject Agency/Client into meta for downstream use
                meta["agency"] = agency_name.strip()
                meta["client_name"] = client_name.strip()

                st.success("Extracted successfully.")
                with st.expander("Meta (preview)", expanded=True):
                    st.write(meta)
                with st.expander("Signposted content (preview)", expanded=True):
                    st.text("\n".join(lines))

                if do_doc:
                    out_bytes = build_docx(tpl_file.read(), meta, lines)
                    fname = safe_filename(f"{meta['page']} - Content Recommendations") + ".docx"
                    # store for stable download across reruns
                    st.session_state.single_docx = out_bytes
                    st.session_state.single_docx_name = fname
            except Exception as e:
                st.exception(e)

    # render download button if we have a generated file
    if st.session_state.single_docx:
        st.download_button(
            "Download DOCX",
            data=st.session_state.single_docx,
            file_name=st.session_state.single_docx_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_single_docx",
        )

with tab2:
    st.subheader("Batch process CSV")
    st.caption("Upload a CSV with a header row; required column: url. Optional: out_name.")
    batch_file = st.file_uploader("CSV file", type=["csv"], key="csv")
    if st.button("Run batch"):
        if not tpl_file:
            st.error("Please upload your Rec Template.docx in the sidebar first.")
        elif not batch_file:
            st.error("Please upload a CSV.")
        else:
            tpl_bytes = tpl_file.read()
            rows = list(csv.DictReader(io.StringIO(batch_file.getvalue().decode("utf-8"))))
            if not rows:
                st.error("CSV appears empty.")
            elif "url" not in rows[0]:
                st.error("CSV must include a 'url' column.")
            else:
                memzip = io.BytesIO()
                zf = zipfile.ZipFile(memzip, "w", zipfile.ZIP_DEFLATED)
                results = []
                for i, row in enumerate(rows, 1):
                    u = row["url"].strip()
                    try:
                        meta, lines = process_url(
                            u,
                            exclude_selectors,
                            annotate_links=annotate_links,
                            remove_before_h1=remove_before_h1,
                            include_img_src=include_img_src,
                        )
                        out_name_raw = (row.get("out_name") or f"{meta['page']} - Content Recommendations").strip()
                        out_name = safe_filename(out_name_raw)
                        out_bytes = build_docx(tpl_bytes, meta, lines)
                        zf.writestr(f"{out_name}.docx", out_bytes)
                        results.append({"url": u, "status": "ok", "file": f"{out_name}.docx"})
                    except Exception as e:
                        results.append({"url": u, "status": f"error: {e}", "file": ""})
                zf.close()
                memzip.seek(0)
                st.success("Batch complete.")
                st.dataframe(results)

                # store for stable download across reruns
                st.session_state.batch_zip = memzip.read()

    # render batch download if available
    if st.session_state.batch_zip:
        st.download_button(
            "Download ZIP",
            data=st.session_state.batch_zip,
            file_name=st.session_state.batch_zip_name,
            mime="application/zip",
            key="dl_batch_zip",
        )
